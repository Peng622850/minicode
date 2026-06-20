from docx import Document

# 创建一个新的 Word 文档
doc = Document()

# 添加标题
doc.add_heading('示例文档', level=1)

# 添加段落
doc.add_paragraph('这是一个新创建的 Word 文档。')
doc.add_paragraph('你可以在这里添加更多的内容。')

# 保存文档
doc.save('example.docx')

print('Word 文档已创建成功！')